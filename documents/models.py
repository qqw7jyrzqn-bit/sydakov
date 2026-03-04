"""
Модели приложения documents.
Управление документами и генерация отчётов.
"""
from django.db import models
from django.conf import settings
from wells.models import Well
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os
from django.core.files import File


class Document(models.Model):
    """
    Модель документа (техзадание, отчёт и т.д.)
    """
    
    DOCUMENT_TYPE_CHOICES = [
        ('technical_spec', 'Техническое задание'),
        ('report', 'Отчёт'),
        ('protocol', 'Протокол'),
        ('geology_report', 'Геологический отчёт'),
        ('drilling_program', 'Программа бурения'),
        ('completion_report', 'Акт завершения'),
        ('safety_plan', 'План по безопасности'),
        ('other', 'Прочее'),
    ]
    
    well = models.ForeignKey(
        Well,
        on_delete=models.CASCADE,
        related_name='documents',
        verbose_name='Скважина'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Название документа'
    )
    document_type = models.CharField(
        max_length=20,
        choices=DOCUMENT_TYPE_CHOICES,
        default='other',
        verbose_name='Тип документа'
    )
    file = models.FileField(
        upload_to='documents/',
        blank=True,
        null=True,
        verbose_name='Файл документа'
    )
    generated_docx = models.FileField(
        upload_to='generated_reports/',
        blank=True,
        null=True,
        verbose_name='Сгенерированный отчёт'
    )
    uploaded_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.PROTECT,
        verbose_name='Загрузил'
    )
    uploaded_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Дата загрузки'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Описание'
    )
    
    class Meta:
        verbose_name = 'Документ'
        verbose_name_plural = 'Документы'
        ordering = ['-uploaded_at']
    
    def __str__(self):
        return f"{self.title} ({self.well.name})"
    
    def generate_docx(self):
        """
        Генерация Word-отчёта по скважине.
        Использует python-docx для создания документа.
        """
        # Создаём новый документ
        doc = DocxDocument()
        
        # Генерируем документ в зависимости от типа
        if self.document_type == 'technical_spec':
            self._generate_technical_spec(doc)
        elif self.document_type == 'protocol':
            self._generate_protocol(doc)
        elif self.document_type == 'report':
            self._generate_report(doc)
        elif self.document_type == 'geology_report':
            self._generate_geology_report(doc)
        elif self.document_type == 'drilling_program':
            self._generate_drilling_program(doc)
        elif self.document_type == 'completion_report':
            self._generate_completion_report(doc)
        elif self.document_type == 'safety_plan':
            self._generate_safety_plan(doc)
        else:
            self._generate_generic_document(doc)
        
        # Сохранение документа
        filename = f'{self.document_type}_well_{self.well.pk}_{self.well.name.replace(" ", "_")}.docx'
        filepath = os.path.join(settings.MEDIA_ROOT, 'generated_reports', filename)
        
        # Создаём директорию если её нет
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        doc.save(filepath)
        
        # Сохраняем файл в модель
        with open(filepath, 'rb') as f:
            # Используем относительный путь для сохранения
            relative_path = os.path.join('generated_reports', filename)
            self.generated_docx.save(relative_path, File(f), save=True)
        
        # Проверяем, что файл действительно сохранён
        if not self.generated_docx or not self.generated_docx.storage.exists(self.generated_docx.name):
            raise Exception(f'Не удалось сохранить файл {filename}')
        
        return self.generated_docx.url
    
    def _apply_document_formatting(self, doc):
        """
        Применяет единое форматирование ко всему документу:
        - Шрифт Times New Roman 14pt
        - Цвет текста черный
        - Межстрочный интервал 1.5
        """
        for paragraph in doc.paragraphs:
            # Устанавливаем межстрочный интервал
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            # Форматируем все run в параграфе
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Форматируем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Устанавливаем межстрочный интервал
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)
                            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _generate_report(self, doc):
        """Генерация отчёта по скважине"""
        # Заголовок
        title = doc.add_heading('ОТЧЁТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'о результатах проектирования и контроля строительства скважины {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о месторождении
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        p = doc.add_paragraph()
        p.add_run(f'Дата составления: ').bold = True
        from datetime import datetime
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        
        doc.add_page_break()
        
        # 1. Общие сведения
        doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ О СКВАЖИНЕ', level=1)
        
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        data = [
            ['Наименование скважины:', self.well.name],
            ['Месторождение:', self.well.field],
            ['Географические координаты:', self.well.coordinates or 'Не указаны'],
            ['Проектная глубина:', f'{self.well.depth} м'],
            ['Текущий статус:', self.well.get_status_display()],
            ['Ответственный инженер:', self.well.author.get_full_name() or self.well.author.username],
            ['Дата начала проектирования:', self.well.created_at.strftime('%d.%m.%Y')],
        ]
        
        for i, (label, value) in enumerate(data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
        
        # 2. Геологическая характеристика
        doc.add_heading('2. ГЕОЛОГИЧЕСКАЯ ХАРАКТЕРИСТИКА', level=1)
        doc.add_paragraph(
            f'Скважина {self.well.name} расположена на {self.well.field} месторождении. '
            f'Проектная глубина бурения составляет {self.well.depth} метров. '
            'Геологический разрез представлен следующими стратиграфическими комплексами:'
        )
        
        doc.add_paragraph('• Четвертичные отложения (0-50м) - суглинки, глины с прослоями песка', style='List Bullet')
        doc.add_paragraph('• Неогеновые отложения (50-350м) - песчано-глинистая толща', style='List Bullet')
        doc.add_paragraph('• Палеогеновые отложения (350-800м) - песчаники, алевролиты, аргиллиты', style='List Bullet')
        doc.add_paragraph(f'• Продуктивный горизонт ({int(float(self.well.depth) * 0.7)}-{self.well.depth}м) - коллекторы с нефтегазонасыщением', style='List Bullet')
        
        # 3. Техническая часть
        doc.add_heading('3. КОНСТРУКЦИЯ СКВАЖИНЫ', level=1)
        doc.add_paragraph(
            'Конструкция скважины предусматривает спуск следующих обсадных колонн:'
        )
        
        construction_table = doc.add_table(rows=4, cols=4)
        construction_table.style = 'Light Grid Accent 1'
        
        hdr_cells = construction_table.rows[0].cells
        hdr_cells[0].text = 'Колонна'
        hdr_cells[1].text = 'Диаметр, мм'
        hdr_cells[2].text = 'Глубина спуска, м'
        hdr_cells[3].text = 'Цементирование'
        
        constructions = [
            ['Направление', '426', '50', 'До устья'],
            ['Кондуктор', '324', '350', 'До устья'],
            ['Эксплуатационная', '168', str(self.well.depth), f'До {int(float(self.well.depth) * 0.5)}м'],
        ]
        
        for i, row_data in enumerate(constructions, 1):
            row_cells = construction_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        
        # 4. Описание работ
        if self.well.description:
            doc.add_heading('4. ОПИСАНИЕ ПРОЕКТНЫХ РЕШЕНИЙ', level=1)
            doc.add_paragraph(self.well.description)
        
        doc.add_paragraph()
        doc.add_paragraph(
            'В процессе бурения предусмотрено применение следующих технологий:'
        )
        doc.add_paragraph('• Роторное бурение с промывкой буровым раствором на водной основе', style='List Bullet')
        doc.add_paragraph('• Использование долот PDC для твердых пород', style='List Bullet')
        doc.add_paragraph('• Контроль параметров бурения в режиме реального времени', style='List Bullet')
        doc.add_paragraph('• Геофизические исследования в процессе бурения (LWD)', style='List Bullet')
        
        # 5. Мероприятия по охране недр
        doc.add_heading('5. МЕРОПРИЯТИЯ ПО ОХРАНЕ НЕДР И ОКРУЖАЮЩЕЙ СРЕДЫ', level=1)
        doc.add_paragraph(
            'При строительстве скважины предусмотрены следующие мероприятия:'
        )
        doc.add_paragraph('• Предотвращение открытых фонтанов и выбросов', style='List Bullet')
        doc.add_paragraph('• Контроль качества цементирования обсадных колонн', style='List Bullet')
        doc.add_paragraph('• Использование противовыбросового оборудования', style='List Bullet')
        doc.add_paragraph('• Утилизация буровых отходов согласно экологическим требованиям', style='List Bullet')
        doc.add_paragraph('• Рекультивация земель после окончания буровых работ', style='List Bullet')
        
        # 6. История согласований
        doc.add_heading('6. ИСТОРИЯ СОГЛАСОВАНИЯ ПРОЕКТА', level=1)
        
        approvals = self.well.approval_steps.all()
        if approvals:
            approval_table = doc.add_table(rows=1, cols=4)
            approval_table.style = 'Light Grid Accent 1'
            
            hdr_cells = approval_table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Должностное лицо'
            hdr_cells[2].text = 'Решение'
            hdr_cells[3].text = 'Комментарий'
            
            for approval in approvals:
                row_cells = approval_table.add_row().cells
                row_cells[0].text = approval.timestamp.strftime('%d.%m.%Y %H:%M')
                row_cells[1].text = approval.user.get_full_name() or approval.user.username
                row_cells[2].text = approval.get_status_display()
                row_cells[3].text = approval.comment or '-'
        else:
            doc.add_paragraph('Согласование проекта в процессе.')
        
        # 7. Заключение
        doc.add_heading('7. ЗАКЛЮЧЕНИЕ', level=1)
        doc.add_paragraph(
            f'Проект строительства скважины {self.well.name} разработан в соответствии '
            'с действующими нормативными документами и правилами безопасности в нефтяной '
            'и газовой промышленности. Предусмотренные технические решения обеспечивают '
            'безопасное и качественное строительство скважины с минимальным воздействием '
            'на окружающую среду.'
        )
        
        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Инженер ПТО: ').bold = True
        p.add_run('_' * 20)
        p.add_run(f'  {self.well.author.get_full_name() or self.well.author.username}')
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_technical_spec(self, doc):
        """Генерация технического задания"""
        # Заголовок
        title = doc.add_heading('ТЕХНИЧЕСКОЕ ЗАДАНИЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'на проектирование и строительство скважины {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        p = doc.add_paragraph()
        p.add_run(f'Утверждено: ').bold = True
        from datetime import datetime
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        
        doc.add_page_break()
        
        # 1. Основание для разработки
        doc.add_heading('1. ОСНОВАНИЕ ДЛЯ РАЗРАБОТКИ', level=1)
        doc.add_paragraph(
            f'Настоящее техническое задание разработано на основании плана развития '
            f'{self.well.field} месторождения и утверждённой программы геологоразведочных работ. '
            'Строительство скважины предусмотрено для уточнения геологического строения, '
            'оценки запасов углеводородов и подготовки месторождения к промышленной эксплуатации.'
        )
        
        # 2. Цель и назначение
        doc.add_heading('2. ЦЕЛЬ И НАЗНАЧЕНИЕ СКВАЖИНЫ', level=1)
        doc.add_paragraph('Основные задачи строительства скважины:')
        doc.add_paragraph('• Вскрытие и опробование продуктивных горизонтов', style='List Bullet')
        doc.add_paragraph('• Получение достоверной геологической информации', style='List Bullet')
        doc.add_paragraph('• Оценка фильтрационно-емкостных свойств коллекторов', style='List Bullet')
        doc.add_paragraph('• Определение физико-химических свойств пластовых флюидов', style='List Bullet')
        doc.add_paragraph('• Уточнение параметров для подсчёта запасов', style='List Bullet')
        
        # 3. Местоположение
        doc.add_heading('3. МЕСТОПОЛОЖЕНИЕ СКВАЖИНЫ', level=1)
        
        location_table = doc.add_table(rows=4, cols=2)
        location_table.style = 'Light Grid Accent 1'
        
        location_data = [
            ['Месторождение:', self.well.field],
            ['Наименование скважины:', self.well.name],
            ['Координаты устья:', self.well.coordinates or 'Определяются на местности'],
            ['Альтитуда:', 'Уточняется по результатам топосъёмки'],
        ]
        
        for i, (label, value) in enumerate(location_data):
            row = location_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
        
        # 4. Геологические условия
        doc.add_heading('4. ГЕОЛОГИЧЕСКИЕ УСЛОВИЯ БУРЕНИЯ', level=1)
        doc.add_paragraph(
            'По данным сейсморазведочных работ и бурения соседних скважин ожидается '
            'следующий геологический разрез:'
        )
        
        geology_table = doc.add_table(rows=5, cols=3)
        geology_table.style = 'Light Grid Accent 1'
        
        hdr_cells = geology_table.rows[0].cells
        hdr_cells[0].text = 'Стратиграфия'
        hdr_cells[1].text = 'Интервал, м'
        hdr_cells[2].text = 'Литология'
        
        geology_data = [
            ['Четвертичные', '0-50', 'Суглинки, глины, пески'],
            ['Неогеновые', '50-350', 'Песчано-глинистая толща'],
            ['Палеогеновые', '350-800', 'Песчаники, алевролиты'],
            [f'Продуктивный', f'{int(float(self.well.depth) * 0.7)}-{self.well.depth}', 'Нефтегазонасыщенные коллекторы'],
        ]
        
        for i, row_data in enumerate(geology_data, 1):
            row_cells = geology_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        
        # 5. Проектная конструкция
        doc.add_heading('5. ПРОЕКТНАЯ КОНСТРУКЦИЯ СКВАЖИНЫ', level=1)
        doc.add_paragraph(f'Проектная глубина скважины: {self.well.depth} м')
        doc.add_paragraph()
        doc.add_paragraph('Конструкция скважины:')
        
        doc.add_paragraph(f'• Направление Ø426мм на глубину 50м', style='List Bullet')
        doc.add_paragraph(f'• Кондуктор Ø324мм на глубину 350м', style='List Bullet')
        doc.add_paragraph(f'• Эксплуатационная колонна Ø168мм на глубину {self.well.depth}м', style='List Bullet')
        
        # 6. Требования к качеству
        doc.add_heading('6. ТРЕБОВАНИЯ К КАЧЕСТВУ ВЫПОЛНЕНИЯ РАБОТ', level=1)
        doc.add_paragraph('При строительстве скважины необходимо обеспечить:')
        doc.add_paragraph('• Качественное цементирование обсадных колонн (коэффициент заполнения не менее 0,8)', style='List Bullet')
        doc.add_paragraph('• Соблюдение проектного профиля скважины (зенитный угол не более 3°)', style='List Bullet')
        doc.add_paragraph('• Предотвращение поглощений и нефтегазоводопроявлений', style='List Bullet')
        doc.add_paragraph('• Качественную проходку продуктивных горизонтов', style='List Bullet')
        doc.add_paragraph('• Соответствие экологическим и санитарным нормам', style='List Bullet')
        
        # 7. Сроки
        doc.add_heading('7. СРОКИ ВЫПОЛНЕНИЯ РАБОТ', level=1)
        doc.add_paragraph(
            f'Проектный срок строительства скважины {self.well.name} составляет 45 календарных дней '
            'с момента начала буровых работ до сдачи скважины заказчику.'
        )
        
        # 8. Требования к документации
        doc.add_heading('8. ТРЕБОВАНИЯ К ДОКУМЕНТАЦИИ', level=1)
        doc.add_paragraph('По результатам строительства скважины подрядчик обязан предоставить:')
        doc.add_paragraph('• Суточные рапорты о ходе буровых работ', style='List Bullet')
        doc.add_paragraph('• Геологический разрез скважины', style='List Bullet')
        doc.add_paragraph('• Результаты геофизических исследований', style='List Bullet')
        doc.add_paragraph('• Протоколы испытаний обсадных колонн', style='List Bullet')
        doc.add_paragraph('• Акт сдачи-приёмки выполненных работ', style='List Bullet')
        
        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Разработал: ').bold = True
        p.add_run('_' * 20)
        p.add_run(f'  {self.well.author.get_full_name() or self.well.author.username}')
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_protocol(self, doc):
        """Генерация протокола испытаний"""
        # Заголовок
        title = doc.add_heading('ПРОТОКОЛ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'испытания и приёмки скважины {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        from datetime import datetime
        p.add_run(f'Дата составления: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        
        doc.add_page_break()
        
        # 1. Общие сведения
        doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ', level=1)
        doc.add_paragraph(
            f'Комиссия в составе представителей Заказчика и Подрядчика произвела осмотр '
            f'и приёмку скважины {self.well.name}, пробуренной на {self.well.field} месторождении.'
        )
        
        # 2. Фактические данные
        doc.add_heading('2. ФАКТИЧЕСКИЕ ПОКАЗАТЕЛИ СТРОИТЕЛЬСТВА', level=1)
        
        facts_table = doc.add_table(rows=6, cols=2)
        facts_table.style = 'Light Grid Accent 1'
        
        facts_data = [
            ['Плановая глубина:', f'{self.well.depth} м'],
            ['Фактическая глубина:', f'{self.well.depth} м'],
            ['Дата начала бурения:', self.well.created_at.strftime('%d.%m.%Y')],
            ['Дата окончания бурения:', datetime.now().strftime('%d.%m.%Y')],
            ['Фактическая продолжительность:', f'{(datetime.now().date() - self.well.created_at.date()).days} сут.'],
            ['Средняя механическая скорость:', f'{float(self.well.depth) / max((datetime.now().date() - self.well.created_at.date()).days, 1):.1f} м/час'],
        ]
        
        for i, (label, value) in enumerate(facts_data):
            row = facts_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
        
        # 3. Конструкция скважины
        doc.add_heading('3. ФАКТИЧЕСКАЯ КОНСТРУКЦИЯ СКВАЖИНЫ', level=1)
        
        construction_table = doc.add_table(rows=4, cols=5)
        construction_table.style = 'Light Grid Accent 1'
        
        hdr_cells = construction_table.rows[0].cells
        hdr_cells[0].text = 'Колонна'
        hdr_cells[1].text = 'Диаметр'
        hdr_cells[2].text = 'Глубина спуска'
        hdr_cells[3].text = 'Цемент'
        hdr_cells[4].text = 'Результат ОЗЦ'
        
        const_data = [
            ['Направление', 'Ø426мм', '50м', 'До устья', 'Положительный'],
            ['Кондуктор', 'Ø324мм', '350м', 'До устья', 'Положительный'],
            ['Эксплуатационная', 'Ø168мм', f'{self.well.depth}м', f'{int(float(self.well.depth) * 0.5)}м', 'Положительный'],
        ]
        
        for i, row_data in enumerate(const_data, 1):
            row_cells = construction_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        
        # 4. Геофизические исследования
        doc.add_heading('4. РЕЗУЛЬТАТЫ ГЕОФИЗИЧЕСКИХ ИССЛЕДОВАНИЙ', level=1)
        doc.add_paragraph('В скважине выполнен следующий комплекс геофизических исследований:')
        doc.add_paragraph('• Стандартный каротаж (ПС, КС, БК)', style='List Bullet')
        doc.add_paragraph('• Радиоактивный каротаж (ГК, НК)', style='List Bullet')
        doc.add_paragraph('• Акустический каротаж (АК)', style='List Bullet')
        doc.add_paragraph('• Боковое каротажное зондирование (БКЗ)', style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph(
            f'По данным ГИС в интервале {int(float(self.well.depth) * 0.75)}-{self.well.depth}м '
            'выделены продуктивные пласты с промышленными запасами углеводородов.'
        )
        
        # 5. Испытания
        doc.add_heading('5. РЕЗУЛЬТАТЫ ИСПЫТАНИЯ ПРОДУКТИВНЫХ ПЛАСТОВ', level=1)
        
        testing_table = doc.add_table(rows=2, cols=4)
        testing_table.style = 'Light Grid Accent 1'
        
        hdr_cells = testing_table.rows[0].cells
        hdr_cells[0].text = 'Интервал, м'
        hdr_cells[1].text = 'Дебит нефти, т/сут'
        hdr_cells[2].text = 'Пластовое давление, МПа'
        hdr_cells[3].text = 'Обводнённость, %'
        
        row_cells = testing_table.rows[1].cells
        row_cells[0].text = f'{int(float(self.well.depth) * 0.85)}-{self.well.depth}'
        row_cells[1].text = '15.3'
        row_cells[2].text = '24.7'
        row_cells[3].text = '2.5'
        
        # 6. Качество работ
        doc.add_heading('6. ОЦЕНКА КАЧЕСТВА ВЫПОЛНЕННЫХ РАБОТ', level=1)
        doc.add_paragraph('Комиссия установила:')
        doc.add_paragraph('• Скважина пробурена в соответствии с утверждённым проектом', style='List Bullet')
        doc.add_paragraph('• Качество цементирования обсадных колонн соответствует требованиям', style='List Bullet')
        doc.add_paragraph('• Профиль скважины выдержан в пределах допустимых отклонений', style='List Bullet')
        doc.add_paragraph('• Продуктивные горизонты вскрыты качественно', style='List Bullet')
        doc.add_paragraph('• Экологические требования соблюдены', style='List Bullet')
        
        # 7. Заключение
        doc.add_heading('7. ЗАКЛЮЧЕНИЕ КОМИССИИ', level=1)
        doc.add_paragraph(
            f'Скважина {self.well.name} построена в соответствии с техническим заданием, '
            'качество выполненных работ соответствует нормативным требованиям. '
            'Комиссия рекомендует скважину к приёмке и вводу в эксплуатацию.'
        )
        
        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('Члены комиссии:')
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('От Заказчика: ').bold = True
        p.add_run('_' * 20)
        p = doc.add_paragraph()
        p.add_run('От Подрядчика: ').bold = True
        p.add_run('_' * 20)
        p.add_run(f'  {self.well.author.get_full_name() or self.well.author.username}')
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_generic_document(self, doc):
        """Генерация общего документа"""
        # Используем базовую генерацию отчёта
        self._generate_report(doc)
    
    def _generate_geology_report(self, doc):
        """Генерация геологического отчёта"""
        from datetime import datetime
        
        title = doc.add_heading('ГЕОЛОГИЧЕСКИЙ ОТЧЁТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'по скважине {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        p = doc.add_paragraph()
        p.add_run(f'Дата: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        
        doc.add_page_break()
        
        doc.add_heading('1. ГЕОЛОГИЧЕСКОЕ СТРОЕНИЕ', level=1)
        doc.add_paragraph(
            f'Скважина {self.well.name} расположена в пределах {self.well.field} месторождения. '
            'Геологический разрез изучен на основании данных сейсморазведки и бурения соседних скважин.'
        )
        
        doc.add_heading('1.1. Стратиграфия', level=2)
        strat_table = doc.add_table(rows=5, cols=4)
        strat_table.style = 'Light Grid Accent 1'
        
        hdr = strat_table.rows[0].cells
        hdr[0].text = 'Система'
        hdr[1].text = 'Интервал, м'
        hdr[2].text = 'Мощность, м'
        hdr[3].text = 'Литология'
        
        strat_data = [
            ['Четвертичная', '0-50', '50', 'Суглинки, глины'],
            ['Неогеновая', '50-350', '300', 'Песчано-глинистая толща'],
            ['Палеогеновая', '350-800', '450', 'Песчаники, алевролиты'],
            [f'Продуктивная', f'{int(float(self.well.depth) * 0.7)}-{self.well.depth}', f'{int(float(self.well.depth) * 0.3)}', 'Коллекторы'],
        ]
        
        for i, row_data in enumerate(strat_data, 1):
            cells = strat_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('1.2. Тектоника', level=2)
        doc.add_paragraph(
            'Территория характеризуется спокойным тектоническим режимом. '
            'Выделяется система разломов северо-восточного простирания.'
        )
        
        doc.add_heading('2. ПЕТРОФИЗИЧЕСКАЯ ХАРАКТЕРИСТИКА', level=1)
        petro_table = doc.add_table(rows=4, cols=3)
        petro_table.style = 'Light Grid Accent 1'
        
        hdr = petro_table.rows[0].cells
        hdr[0].text = 'Параметр'
        hdr[1].text = 'Значение'
        hdr[2].text = 'Ед. изм.'
        
        petro_data = [
            ['Пористость', '18-22', '%'],
            ['Проницаемость', '150-350', 'мД'],
            ['Нефтенасыщенность', '65-75', '%'],
        ]
        
        for i, row_data in enumerate(petro_data, 1):
            cells = petro_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('3. ГИДРОГЕОЛОГИЯ', level=1)
        doc.add_paragraph(
            'Водоносные горизонты приурочены к четвертичным и неогеновым отложениям. '
            'Уровень грунтовых вод находится на глубине 15-25 метров.'
        )
        
        doc.add_heading('4. ВЫВОДЫ', level=1)
        doc.add_paragraph(
            f'По результатам геологических исследований скважина {self.well.name} '
            'может рассматриваться как перспективная для промышленной эксплуатации.'
        )
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Геолог: ').bold = True
        p.add_run('_' * 20)
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_drilling_program(self, doc):
        """Генерация программы бурения"""
        from datetime import datetime
        
        title = doc.add_heading('ПРОГРАММА БУРЕНИЯ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'скважины {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        p = doc.add_paragraph()
        p.add_run(f'Утверждена: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        
        doc.add_page_break()
        
        doc.add_heading('1. ОБЩИЕ ПОЛОЖЕНИЯ', level=1)
        doc.add_paragraph(
            f'Настоящая программа разработана для бурения скважины {self.well.name} '
            f'на {self.well.field} месторождении. Проектная глубина {self.well.depth} м.'
        )
        
        doc.add_heading('2. РЕЖИМЫ БУРЕНИЯ', level=1)
        regime_table = doc.add_table(rows=4, cols=4)
        regime_table.style = 'Light Grid Accent 1'
        
        hdr = regime_table.rows[0].cells
        hdr[0].text = 'Интервал, м'
        hdr[1].text = 'Нагрузка, т'
        hdr[2].text = 'Обороты, об/мин'
        hdr[3].text = 'Расход, л/с'
        
        regime_data = [
            ['0-50', '8-10', '80-100', '25-30'],
            ['50-350', '10-12', '70-90', '30-35'],
            [f'350-{self.well.depth}', '12-15', '60-80', '35-40'],
        ]
        
        for i, row_data in enumerate(regime_data, 1):
            cells = regime_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('3. БУРОВЫЕ РАСТВОРЫ', level=1)
        doc.add_paragraph('Рекомендуемые параметры буровых растворов:')
        doc.add_paragraph('• Плотность: 1.08-1.12 г/см³', style='List Bullet')
        doc.add_paragraph('• Вязкость: 30-40 сек', style='List Bullet')
        doc.add_paragraph('• Водоотдача: не более 8 см³/30мин', style='List Bullet')
        doc.add_paragraph('• pH: 9-10', style='List Bullet')
        
        doc.add_heading('4. ДОЛОТА', level=1)
        dolota_table = doc.add_table(rows=4, cols=3)
        dolota_table.style = 'Light Grid Accent 1'
        
        hdr = dolota_table.rows[0].cells
        hdr[0].text = 'Интервал, м'
        hdr[1].text = 'Тип долота'
        hdr[2].text = 'Диаметр, мм'
        
        dolota_data = [
            ['0-50', 'Шарошечное', '490'],
            ['50-350', 'PDC', '393.7'],
            [f'350-{self.well.depth}', 'PDC', '215.9'],
        ]
        
        for i, row_data in enumerate(dolota_data, 1):
            cells = dolota_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('5. ОБСАДНЫЕ КОЛОННЫ', level=1)
        doc.add_paragraph(f'• Направление Ø426мм на глубину 50м', style='List Bullet')
        doc.add_paragraph(f'• Кондуктор Ø324мм на глубину 350м', style='List Bullet')
        doc.add_paragraph(f'• Эксплуатационная Ø168мм на глубину {self.well.depth}м', style='List Bullet')
        
        doc.add_heading('6. КОНТРОЛЬ КАЧЕСТВА', level=1)
        doc.add_paragraph('В процессе бурения осуществлять:')
        doc.add_paragraph('• Ежесуточный контроль параметров бурового раствора', style='List Bullet')
        doc.add_paragraph('• Геофизические исследования в процессе бурения', style='List Bullet')
        doc.add_paragraph('• Отбор керна в продуктивных интервалах', style='List Bullet')
        doc.add_paragraph('• Контроль отклонения ствола скважины', style='List Bullet')
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Главный инженер: ').bold = True
        p.add_run('_' * 20)
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_completion_report(self, doc):
        """Генерация акта завершения"""
        from datetime import datetime
        
        title = doc.add_heading('АКТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'о завершении строительства скважины {self.well.name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Дата: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        p = doc.add_paragraph()
        p.add_run(f'Месторождение: ').bold = True
        p.add_run(self.well.field)
        
        doc.add_page_break()
        
        doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ', level=1)
        doc.add_paragraph(
            f'Комиссия в составе представителей Заказчика и Подрядчика составила настоящий акт '
            f'о завершении строительства скважины {self.well.name}.'
        )
        
        doc.add_heading('2. ВЫПОЛНЕННЫЕ РАБОТЫ', level=1)
        works_table = doc.add_table(rows=6, cols=2)
        works_table.style = 'Light Grid Accent 1'
        
        hdr = works_table.rows[0].cells
        hdr[0].text = 'Вид работ'
        hdr[1].text = 'Статус'
        
        works_data = [
            ['Бурение скважины', 'Выполнено'],
            ['Спуск обсадных колонн', 'Выполнено'],
            ['Цементирование', 'Выполнено'],
            ['Геофизические исследования', 'Выполнено'],
            ['Испытания пластов', 'Выполнено'],
        ]
        
        for i, row_data in enumerate(works_data, 1):
            cells = works_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('3. ТЕХНИЧЕСКИЕ ПОКАЗАТЕЛИ', level=1)
        tech_table = doc.add_table(rows=5, cols=2)
        tech_table.style = 'Light Grid Accent 1'
        
        tech_data = [
            ['Плановая глубина:', f'{self.well.depth} м'],
            ['Фактическая глубина:', f'{self.well.depth} м'],
            ['Продолжительность строительства:', f'{(datetime.now().date() - self.well.created_at.date()).days} сут.'],
            ['Средняя скорость бурения:', f'{float(self.well.depth) / max((datetime.now().date() - self.well.created_at.date()).days, 1):.1f} м/сут'],
            ['Качество цементирования:', 'Удовлетворительное'],
        ]
        
        for i, (label, value) in enumerate(tech_data):
            cells = tech_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
        
        doc.add_heading('4. ЗАКЛЮЧЕНИЕ', level=1)
        doc.add_paragraph(
            f'Строительство скважины {self.well.name} завершено в соответствии с проектом. '
            'Скважина готова к вводу в эксплуатацию.'
        )
        
        doc.add_paragraph()
        doc.add_paragraph('Подписи членов комиссии:')
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('От Заказчика: ').bold = True
        p.add_run('_' * 25)
        p = doc.add_paragraph()
        p.add_run('От Подрядчика: ').bold = True
        p.add_run('_' * 25)
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
    
    def _generate_safety_plan(self, doc):
        """Генерация плана по безопасности"""
        from datetime import datetime
        
        title = doc.add_heading('ПЛАН МЕРОПРИЯТИЙ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('по промышленной безопасности и охране труда', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle2 = doc.add_heading(f'при строительстве скважины {self.well.name}', level=2)
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Утверждён: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y'))
        
        doc.add_page_break()
        
        doc.add_heading('1. ОБЩИЕ ТРЕБОВАНИЯ БЕЗОПАСНОСТИ', level=1)
        doc.add_paragraph(
            'К работам допускаются лица, прошедшие медицинский осмотр, обучение и инструктаж '
            'по охране труда и промышленной безопасности.'
        )
        
        doc.add_heading('2. ОРГАНИЗАЦИОННЫЕ МЕРОПРИЯТИЯ', level=1)
        doc.add_paragraph('• Назначение ответственных лиц', style='List Bullet')
        doc.add_paragraph('• Проведение инструктажей', style='List Bullet')
        doc.add_paragraph('• Обеспечение СИЗ', style='List Bullet')
        doc.add_paragraph('• Контроль состояния оборудования', style='List Bullet')
        doc.add_paragraph('• Организация аварийно-спасательных служб', style='List Bullet')
        
        doc.add_heading('3. ТЕХНИЧЕСКИЕ МЕРОПРИЯТИЯ', level=1)
        measures_table = doc.add_table(rows=7, cols=3)
        measures_table.style = 'Light Grid Accent 1'
        
        hdr = measures_table.rows[0].cells
        hdr[0].text = 'Мероприятие'
        hdr[1].text = 'Ответственный'
        hdr[2].text = 'Срок'
        
        measures_data = [
            ['Проверка ПВО', 'Буровой мастер', 'Ежедневно'],
            ['Испытание манифольда', 'Механик', 'Еженедельно'],
            ['Контроль газопоказаний', 'Буровая бригада', 'Постоянно'],
            ['Проверка систем сигнализации', 'Электрик', 'Еженедельно'],
            ['Учения по ликвидации аварий', 'Начальник буровой', 'Ежемесячно'],
            ['Метеонаблюдения', 'Дежурный персонал', 'Постоянно'],
        ]
        
        for i, row_data in enumerate(measures_data, 1):
            cells = measures_table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                cells[j].text = cell_data
        
        doc.add_heading('4. ДЕЙСТВИЯ ПРИ АВАРИЙНЫХ СИТУАЦИЯХ', level=1)
        doc.add_paragraph('4.1. При газонефтеводопроявлении:', style='Heading 3')
        doc.add_paragraph('• Немедленно закрыть превенторы', style='List Bullet')
        doc.add_paragraph('• Оповестить руководство', style='List Bullet')
        doc.add_paragraph('• Контролировать давление', style='List Bullet')
        
        doc.add_paragraph('4.2. При пожаре:', style='Heading 3')
        doc.add_paragraph('• Включить систему пожаротушения', style='List Bullet')
        doc.add_paragraph('• Эвакуировать персонал', style='List Bullet')
        doc.add_paragraph('• Вызвать пожарную службу', style='List Bullet')
        
        doc.add_heading('5. СРЕДСТВА ИНДИВИДУАЛЬНОЙ ЗАЩИТЫ', level=1)
        doc.add_paragraph('• Защитные каски', style='List Bullet')
        doc.add_paragraph('• Спецодежда', style='List Bullet')
        doc.add_paragraph('• Защитные очки', style='List Bullet')
        doc.add_paragraph('• Защитная обувь', style='List Bullet')
        doc.add_paragraph('• Средства защиты органов дыхания', style='List Bullet')
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Ответственный за ОТ и ПБ: ').bold = True
        p.add_run('_' * 20)
        
        # Применяем форматирование
        self._apply_document_formatting(doc)
